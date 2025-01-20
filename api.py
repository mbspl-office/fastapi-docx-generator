import os
import random
import string
from fastapi import FastAPI, UploadFile, Form
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import Dict
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches
from io import BytesIO
from pathlib import Path
import traceback

app = FastAPI()

# Enable CORS for development purposes
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directory to store images (ensure these directories are created in your project structure)
images_dir = Path("images")
images_dir.mkdir(exist_ok=True)

# Helper function to generate a random 7-character string for filenames
def generate_random_filename(length=7):
    return ''.join(random.choices(string.ascii_letters + string.digits, k=length))

class Payload(BaseModel):
    data: Dict[str, str]

@app.post("/generate-docx/")
async def generate_docx():
    try:
        # Log the beginning of the docx generation
        print("Starting DOCX generation process...")

        # Load the template DOCX file
        template_path = os.path.join(os.path.dirname(__file__), "SBI Format.docx")
        print(f"Template path: {template_path}")
        
        if not os.path.exists(template_path):
            raise FileNotFoundError("Template file 'SBI Format.docx' not found.")
        
        # Open the template
        doc = Document(template_path)

        # Modify the tables with specific content (same as original code)
        table = doc.tables[0]
        table_input_position = [5, 5, 5, 5, 5, 5, 5, 5, 5]

        for i in range(len(table.rows)):
            if i <= 8:
                table.cell(i, table_input_position[i]).text = "  hello"
            elif i == 23:
                table.cell(i, 5).text = "  hello"

        # Access the second table (index 1)
        table_2 = doc.tables[1]
        table_input_position_2 = [2, 2, 2]
        for i in range(len(table_2.rows)):
            if 12 <= i < 15:
                table_2.cell(i, table_input_position_2[i - 12]).text = "  hello"
            elif i == 16:
                table_2.cell(i, 2).text = "  hello"
            elif i == 18:
                table_2.cell(i, 2).text = "  hello"

        # Access the third table (index 2)
        table_3 = doc.tables[2]
        for row in table_3.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.clear()

        # Add images to the third table
        image_index = 0
        start_row = 0
        start_col = 0
        image_paths = [
            str(images_dir / "1E2KS.jpg"),
            str(images_dir / "1X0I8.jpg"),
            str(images_dir / "2A1VW.jpg"),
            str(images_dir / "2GW4H.jpg")
        ]

        for row_idx in range(start_row, len(table_3.rows)):
            for col_idx in range(start_col, len(table_3.columns)):
                if image_index >= len(image_paths):
                    break  # Stop if no more images to add
                
                cell = table_3.cell(row_idx, col_idx)
                paragraph = cell.paragraphs[0]  # Get the first paragraph in the cell
                run = paragraph.add_run()
                try:
                    run.add_picture(image_paths[image_index], width=Inches(4.06))  # Adjust width as needed
                except Exception as e:
                    print(f"Error adding image {image_paths[image_index]}: {e}")
                image_index += 1  # Move to the next image
            if image_index >= len(image_paths):
                break

        # Generate a random 7-character string for the filename
        random_filename = generate_random_filename() + ".docx"
        
        # Save the document to a BytesIO object instead of a temporary file
        doc_buffer = BytesIO()
        doc.save(doc_buffer)

        # Move the cursor to the beginning of the BytesIO buffer
        doc_buffer.seek(0)

        # Return the generated file for download directly (without saving it to disk)
        return FileResponse(doc_buffer, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=random_filename)

    except Exception as e:
        # Log the full exception and traceback for better diagnosis
        print(f"Error occurred: {str(e)}")
        traceback.print_exc()
        return JSONResponse({"error": f"Internal Server Error: {str(e)}"}, status_code=500)
